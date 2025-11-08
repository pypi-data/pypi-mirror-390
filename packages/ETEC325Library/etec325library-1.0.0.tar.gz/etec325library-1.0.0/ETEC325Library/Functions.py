import numpy as np
import pandas as pd
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
import matplotlib.pyplot as plt

stdoutInstance = sys.stdout

class _WordDocument():
    def __init__(self, file):
        self.file = file + ".docx";
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        style = self.document.styles['No Spacing']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        self.text = ''

    def writeText(self, text):
        if self.text == '':
            self.text = text
        else:
            self.text = self.text + text

    def saveDocument(self):
        paragraph = self.document.add_paragraph(self.text)
        paragraph.style = self.document.styles['No Spacing']
        self.document.save(self.file)

class _Logger(object):
    def __init__(self, file):
        self.terminal = sys.stdout
        self.wordDocument = _WordDocument(file)
   
    def write(self, message):
        self.terminal.write(message)
        self.wordDocument.writeText(message)

    def closeLog(self):
        self.wordDocument.saveDocument()

    def flush(self):
        # this flush method is needed for python 3 compatibility.
        # this handles the flush command by doing nothing.
        # you might want to specify some extra behavior here.
        pass

    def __del__(self):
        self.closeLog()
        sys.stdout = stdoutInstance

def enableLogging(file):
    sys.stdout = _Logger(file)

def disableLogging(file):
    sys.stdout.closeLog()
    sys.stdout = stdoutInstance

def createRange(start, stop, step=1):
    inclusiveRange = np.arange(start, stop + step, step)
    return inclusiveRange
    
def printTable(array, columnNames=None):
    arrayShape = array.shape
    if len(arrayShape) == 1:
        array = np.reshape(array, (1, arrayShape[0]))
        arrayShape = array.shape
    if columnNames is None:
        if len(arrayShape) == 1:
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            df = pd.DataFrame(array, index=blankRows, columns=[''])
        elif len(arrayShape) > 1:
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            numColumns = int(arrayShape[1])
            blankColumns = ['' for i in range(numColumns)]
            df = pd.DataFrame(array, index=blankRows, columns=blankColumns)
    else:
        numRows = int(arrayShape[0])
        blankRows = ['' for i in range(numRows)]
        df = pd.DataFrame(array, index=blankRows, columns=columnNames)
    pd.set_option("display.max_rows", None)
    pd.set_option("display.max_columns", None)
    pd.set_option('display.expand_frame_repr', False)
    print(df)
    print("")
    
def printVariable(label, var, columnNames=None):
    if type(var) is np.ndarray:
        print(label + ":")
        printTable(var, columnNames)
    elif type(var) is list:
        print(label + ":")
        for item in var:
            print("\t", item)
        print("")
    else:
        print(label + ": " + str(var))

def readExcelSpreadsheet(file, useColumnNames=True, useRowLabels=False):
    if useColumnNames:
        headerIndex = 0
    else:
        headerIndex = None

    if useRowLabels:
        rowIndex = 0
    else:
        rowIndex = None
    data = pd.read_excel(file, header=headerIndex, index_col=rowIndex)
    dataNumbers = data.to_numpy()
    indexList = list(data.index)
    namesList = list(data.columns)

    return dataNumbers, indexList, namesList

def getMLDataFromExcelSpreadsheet(file, useColumnNames=True, categoriesCol=-1):
    if useColumnNames:
        headerIndex = 0
    else:
        headerIndex = None

    data = pd.read_excel(file, header=headerIndex, index_col=categoriesCol)
    X = data.to_numpy()
    Y = None
    indexList = list(data.index)
    uniqueIndexesList = list(data.index.unique())
    numClasses = len(uniqueIndexesList)
    firstCategory = indexList[0]
    try:
        int(firstCategory)
        YIdx = [uniqueIndexesList.index(idx) for idx in uniqueIndexesList]
        YData = [int(idx) for idx in uniqueIndexesList]
        YDiff = set(YIdx).difference(YData)
        if len(YDiff) == 0:
            Y = np.array([int(idx) for idx in indexList])
        else:
            Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
    except ValueError:
        Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
    namesList = list(data.columns)

    return X, Y, uniqueIndexesList, numClasses, namesList

def getRawTextMLDataFromExcelSpreadsheet(file, useColumnNames=True, categoriesCol=-1):
    if useColumnNames:
        headerIndex = 0
    else:
        headerIndex = None

    data = pd.read_excel(file, header=headerIndex, index_col=categoriesCol)
    XRaw = data[data.columns[0]].tolist()
    XRaw = [str(s) for s in XRaw]
    XRaw = [s.replace("\\n", "\n") for s in XRaw]
    XRaw = [s.replace("\\r", "\r") for s in XRaw]
    XRaw = [s.replace("\\t", "\t") for s in XRaw]
    Y = None
    indexList = list(data.index)
    uniqueIndexesList = list(data.index.unique())
    numClasses = len(uniqueIndexesList)
    firstCategory = indexList[0]
    try:
        int(firstCategory)
        YIdx = [uniqueIndexesList.index(idx) for idx in uniqueIndexesList]
        YData = [int(idx) for idx in uniqueIndexesList]
        YDiff = set(YIdx).difference(YData)
        if len(YDiff) == 0:
            Y = np.array([int(idx) for idx in indexList])
        else:
            Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
    except ValueError:
        Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
    namesList = list(data.columns)

    return XRaw, Y, uniqueIndexesList, numClasses, namesList
