"""URL fetching utility with Chrome-like user agent."""

import logging
import requests
from typing import Dict, Optional, Any, List, cast
import chardet

from typing import override
from urllib.parse import urlparse
import time
import io
from types import TracebackType


# Text extraction imports
from bs4 import BeautifulSoup
from readability import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


logger = logging.getLogger(__name__)


def _detect_encoding(
    content_bytes: bytes, content_type: str, declared_encoding: Optional[str]
) -> str:
    """Detect the most appropriate encoding for content bytes.

    Args:
        content_bytes: Raw content bytes
        content_type: Content-Type header value
        declared_encoding: Encoding declared in HTTP headers or HTML meta tags

    Returns:
        The most appropriate encoding to use
    """
    # Step 1: Try UTF-8 first since it's the most common web encoding
    try:
        content_bytes.decode("utf-8")
        return "utf-8"
    except UnicodeDecodeError:
        pass

    # Step 2: Try declared encoding from HTTP headers if available and reliable
    if declared_encoding and declared_encoding.lower() not in ["iso-8859-1", "ascii"]:
        try:
            content_bytes.decode(declared_encoding)
            return declared_encoding
        except (UnicodeDecodeError, LookupError):
            logger.debug(
                f"Declared encoding '{declared_encoding}' failed, trying detection"
            )

    # Step 3: For HTML content, try to find meta charset
    if "html" in content_type.lower():
        try:
            # Look for charset in first 1024 bytes (typical location for meta tags)
            head_sample = content_bytes[:1024].decode("ascii", errors="ignore").lower()

            # Look for <meta charset="..."> or <meta http-equiv="content-type" content="...charset=...">
            import re

            charset_patterns = [
                r'<meta\s+charset=["\']?([^"\'>\s]+)',
                r'<meta\s+http-equiv=["\']?content-type["\']?\s+content=["\'][^"\']*charset=([^"\'>\s;]+)',
                r'charset=([^"\'>\s;]+)',
            ]

            for pattern in charset_patterns:
                match = re.search(pattern, head_sample)
                if match:
                    detected_charset = match.group(1).strip()
                    try:
                        content_bytes.decode(detected_charset)
                        return detected_charset
                    except (UnicodeDecodeError, LookupError):
                        continue

        except Exception:
            pass  # Continue to chardet if meta parsing fails

    # Step 4: Use chardet for automatic detection
    try:
        # For performance, only analyze first 10KB for large content
        sample_size = min(len(content_bytes), 10240)
        detection = chardet.detect(content_bytes[:sample_size])

        if detection and detection.get("encoding"):
            detected_encoding = detection["encoding"]
            confidence = detection.get("confidence", 0)

            # Only use chardet result if confidence is reasonable and encoding is not None
            if confidence > 0.7 and detected_encoding:
                try:
                    content_bytes.decode(detected_encoding)
                    return detected_encoding
                except (UnicodeDecodeError, LookupError):
                    pass

    except Exception as e:
        logger.debug(f"Chardet detection failed: {e}")

    # Step 5: Common encoding fallbacks
    fallback_encodings = ["utf-8", "utf-8-sig", "latin1", "cp1252"]

    for encoding in fallback_encodings:
        try:
            content_bytes.decode(encoding)
            return encoding
        except UnicodeDecodeError:
            continue

    # Step 6: Final fallback - use utf-8 with error handling
    return "utf-8"


class URLFetcher:
    """Utility class for fetching web content with Chrome-like behavior."""

    DEFAULT_USER_AGENT = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36"
    )
    """Chrome-like user agent string (Chrome 121 on Windows 10)."""

    DEFAULT_HEADERS = {
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "Pragma": "no-cache",
        "Sec-Ch-Ua": '"Not A(Brand";v="99", "Google Chrome";v="121", "Chromium";v="121"',
        "Sec-Ch-Ua-Mobile": "?0",
        "Sec-Ch-Ua-Platform": '"Windows"',
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "none",
        "Sec-Fetch-User": "?1",
        "Upgrade-Insecure-Requests": "1",
    }
    """Default headers that mimic Chrome browser."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize URL fetcher with optional configuration.

        Args:
            config: Optional configuration dictionary with options:

                - timeout: Request timeout in seconds (default: 30)
                - user_agent: Custom user agent string
                - headers: Additional headers to include
                - follow_redirects: Whether to follow redirects (default: True)
                - max_redirects: Maximum number of redirects (default: 10)

        """
        self.config = config or {}
        self.timeout = self.config.get("timeout", 30)
        self.follow_redirects = self.config.get("follow_redirects", True)
        self.max_redirects = self.config.get("max_redirects", 10)

        # Build headers
        self.headers = self.DEFAULT_HEADERS.copy()
        self.headers["User-Agent"] = self.config.get(
            "user_agent", self.DEFAULT_USER_AGENT
        )

        # Add any custom headers
        custom_headers = self.config.get("headers", {})
        if custom_headers:
            self.headers.update(custom_headers)

        # Create session for connection reuse
        self.session = requests.Session()
        self.session.headers.update(self.headers)

    def fetch(self, url: str, **kwargs: Any) -> Dict[str, Any]:
        """Fetch content from a URL.

        Args:
            url: URL to fetch
            **kwargs: Additional arguments to pass to requests.get()

        Returns:
            Dictionary containing:
                - url: Final URL after redirects
                - status_code: HTTP status code
                - headers: Response headers
                - content: Response content as text
                - content_type: Content type from headers
                - encoding: Character encoding
                - size: Content size in bytes
                - response_time: Time taken for request in seconds

        Raises:
            ValueError: If URL is invalid
            RuntimeError: If request fails
        """
        if not url or not url.strip():
            raise ValueError("URL cannot be empty")

        # Validate URL
        parsed = urlparse(url)
        if not parsed.scheme or not parsed.netloc:
            raise ValueError(f"Invalid URL format: {url}")

        logger.info(f"Fetching URL: {url}")
        start_time = time.time()

        try:
            # Merge kwargs with defaults
            request_kwargs = {
                "timeout": self.timeout,
                "allow_redirects": self.follow_redirects,
            }

            # Note: requests doesn't have max_redirects parameter, it uses DEFAULT_REDIRECT_LIMIT
            # We store max_redirects for potential future use but don't pass it to requests

            request_kwargs.update(kwargs)

            # Make the request
            response = self.session.get(url, **request_kwargs)
            response_time = time.time() - start_time

            # Log response info
            logger.info(f"Response: {response.status_code} in {response_time:.2f}s")

            # Check for HTTP errors
            response.raise_for_status()

            # Extract content information
            content_bytes = response.content
            content_type = (
                response.headers.get("content-type", "").split(";")[0].strip()
            )
            # Content should now be properly decompressed by requests library
            # since we only request gzip/deflate which it handles reliably

            # Use robust encoding detection
            declared_encoding = response.encoding
            detected_encoding = _detect_encoding(
                content_bytes, content_type, declared_encoding
            )

            # Decode content using the detected encoding
            try:
                content = content_bytes.decode(detected_encoding)
            except UnicodeDecodeError:
                # Final fallback with error replacement
                content = content_bytes.decode(detected_encoding, errors="replace")
                logger.warning(
                    f"Used error replacement during decoding with {detected_encoding}"
                )

            encoding = detected_encoding
            size = len(content_bytes)

            result: Dict[str, Any] = {
                "url": response.url,  # Final URL after redirects
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "content": content,
                "content_bytes": content_bytes,
                "content_type": content_type,
                "encoding": encoding,
                "size": size,
                "response_time": response_time,
            }

            logger.info(f"Successfully fetched {size} bytes ({content_type})")
            return result

        except requests.exceptions.Timeout:
            error_msg = f"URLFetcher fetch operation failed: Request timed out after {self.timeout} seconds"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except requests.exceptions.ConnectionError as e:
            error_msg = f"URLFetcher fetch operation failed: Failed to connect to '{url}': {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except requests.exceptions.HTTPError as e:
            # response is available here since HTTPError is raised after response
            response_code = (
                getattr(e.response, "status_code", "unknown")
                if hasattr(e, "response")
                else "unknown"
            )
            error_msg = f"URLFetcher fetch operation failed: HTTP error {response_code} for '{url}': {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except requests.exceptions.RequestException as e:
            error_msg = f"URLFetcher fetch operation failed: Request to '{url}' failed: {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as e:
            error_msg = f"URLFetcher fetch operation failed: Unexpected error fetching '{url}': {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)

    def fetch_headers_only(self, url: str, **kwargs: Any) -> Dict[str, Any]:
        """Fetch only headers from a URL using HEAD request.

        Args:
            url: URL to fetch headers from
            **kwargs: Additional arguments to pass to requests.head()

        Returns:
            Dictionary containing:
                - url: Final URL after redirects
                - status_code: HTTP status code
                - headers: Response headers
                - content_type: Content type from headers
                - content_length: Content length if available
                - response_time: Time taken for request in seconds
        """
        if not url or not url.strip():
            raise ValueError("URL cannot be empty")

        logger.info(f"Fetching headers for URL: {url}")
        start_time = time.time()

        try:
            request_kwargs = {
                "timeout": self.timeout,
                "allow_redirects": self.follow_redirects,
            }
            request_kwargs.update(kwargs)

            response = self.session.head(url, **request_kwargs)
            response_time = time.time() - start_time

            logger.info(
                f"Headers response: {response.status_code} in {response_time:.2f}s"
            )
            response.raise_for_status()

            content_type = (
                response.headers.get("content-type", "").split(";")[0].strip()
            )
            content_length_str = response.headers.get("content-length")
            content_length: Optional[int] = None
            if content_length_str:
                content_length = int(content_length_str)

            return {
                "url": response.url,
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "content_type": content_type,
                "content_length": content_length,
                "response_time": response_time,
            }

        except requests.exceptions.RequestException as e:
            error_msg = f"URLFetcher fetch_headers_only operation failed: Request to '{url}' failed: {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)

    def close(self) -> None:
        """Close the session to free up resources."""
        if hasattr(self, "session"):
            self.session.close()

    def __enter__(self) -> "URLFetcher":
        """Context manager entry."""
        return self

    def __exit__(
        self,
        exc_type: Optional[type],
        exc_val: Optional[BaseException],
        exc_tb: Optional[TracebackType],
    ) -> None:
        """Context manager exit."""
        del exc_type, exc_val, exc_tb  # Suppress pylance warnings
        self.close()

    @override
    def __str__(self) -> str:
        """String representation."""
        return f"URLFetcher(timeout={self.timeout}s)"


# Convenience function for simple URL fetching
def fetch_url(
    url: str, config: Optional[Dict[str, Any]] = None, **kwargs: Any
) -> Dict[str, Any]:
    """Convenience function to fetch a URL with Chrome-like headers.

    Args:
        url: URL to fetch
        config: Optional configuration for URLFetcher
        **kwargs: Additional arguments to pass to the fetch method

    Returns:
        Dictionary with fetch results (see URLFetcher.fetch for details)

    Example:

    .. code-block:: python

        result = fetch_url("https://example.com")
        print(result["content"])
        print(f"Status: {result['status_code']}")

    """
    with URLFetcher(config) as fetcher:
        return fetcher.fetch(url, **kwargs)


def fetch_url_headers(
    url: str, config: Optional[Dict[str, Any]] = None, **kwargs: Any
) -> Dict[str, Any]:
    """Convenience function to fetch only headers from a URL.

    Args:
        url: URL to fetch headers from
        config: Optional configuration for URLFetcher
        **kwargs: Additional arguments to pass to the fetch_headers_only method

    Returns:
        Dictionary with header results (see URLFetcher.fetch_headers_only for details)

    Example:
        >>> headers = fetch_url_headers("https://example.com")
        >>> print(headers["content_type"])
        >>> print(f"Content-Length: {headers['content_length']}")
    """
    with URLFetcher(config) as fetcher:
        return fetcher.fetch_headers_only(url, **kwargs)


def extract_html_text(content: str, method: str = "readability") -> str:
    """Extract main text content from HTML.

    Args:
        content: HTML content as string
        method: Extraction method ("readability" or "basic")

    Returns:
        Extracted text content
    """

    if method == "readability":
        try:
            # Temporarily redirect stderr to suppress readability's error output
            import sys
            from io import StringIO

            # Save original stderr
            original_stderr = sys.stderr

            try:
                # Redirect stderr to a null device
                sys.stderr = StringIO()

                # Use readability to extract main article content
                doc = Document(content)
                article_html = cast(str, doc.summary())

                # Parse with BeautifulSoup to get clean text
                soup = BeautifulSoup(article_html, "html.parser")

                # Remove unwanted elements
                for element in soup(
                    ["script", "style", "nav", "footer", "header", "aside"]
                ):
                    element.decompose()

                # Get text with some formatting preserved
                text = soup.get_text(separator="\n", strip=True)

                # Clean up excessive whitespace
                lines = [line.strip() for line in text.split("\n") if line.strip()]
                return "\n".join(lines)

            finally:
                # Restore original stderr
                sys.stderr = original_stderr

        except Exception as e:
            logger.warning(
                f"Readability extraction failed, falling back to basic: {type(e).__name__}"
            )
            method = "basic"

    if method == "basic":
        # Basic text extraction using BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")

        # Remove unwanted elements
        for element in soup(
            ["script", "style", "nav", "footer", "header", "aside", "form"]
        ):
            element.decompose()

        # Try to find main content areas first
        main_content = None
        for selector in [
            "article",
            "main",
            '[role="main"]',
            ".content",
            "#content",
            ".post",
            ".article",
        ]:
            main_content = soup.select_one(selector)  # pyright: ignore
            if main_content:
                break

        if main_content:
            text = main_content.get_text(separator="\n", strip=True)
        else:
            # Fall back to body content
            body = soup.find("body")
            if body:
                text = body.get_text(separator="\n", strip=True)
            else:
                text = soup.get_text(separator="\n", strip=True)

        # Clean up excessive whitespace
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        return "\n".join(lines)

    raise ValueError(f"Unsupported extraction method: {method}")


def extract_pdf_text(content: bytes) -> str:
    """Extract text content from PDF bytes.

    Args:
        content: PDF content as bytes

    Returns:
        Extracted text content
    """

    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PdfReader(pdf_file)

        text_parts: List[str] = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---")
                    text_parts.append(page_text.strip())
                    text_parts.append("")  # Empty line between pages
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")

        if not text_parts:
            raise RuntimeError("No text content could be extracted from PDF")

        return "\n".join(text_parts).strip()

    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {e}")


def extract_docx_text(content: bytes) -> str:
    """Extract text content from DOCX bytes.

    Args:
        content: DOCX content as bytes

    Returns:
        Extracted text content
    """

    try:
        docx_file = io.BytesIO(content)
        doc = DocxDocument(docx_file)

        text_parts: List[str] = []

        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text: List[str] = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if not text_parts:
            raise RuntimeError("No text content could be extracted from DOCX")

        return "\n".join(text_parts)

    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")


def fetch_main_text(
    url: str, config: Optional[Dict[str, Any]] = None, **kwargs: Any
) -> Dict[str, Any]:
    """Fetch and extract main text content from a URL.

    This function automatically detects content type and extracts readable text:
    - HTML pages: Extracts main article content using readability algorithm
    - PDF files: Extracts text from all pages
    - DOCX files: Extracts text from paragraphs and tables
    - Plain text: Returns as-is

    Args:
        url: URL to fetch and extract text from
        config: Optional configuration for URLFetcher, plus text extraction options:
            - html_method: "readability" (default) or "basic" for HTML extraction
            - include_metadata: Whether to include extraction metadata (default: True)
        **kwargs: Additional arguments to pass to fetch_url

    Returns:
        Dictionary containing:
            - url: Final URL after redirects
            - content_type: Detected content type
            - text: Extracted text content
            - extraction_method: Method used for text extraction
            - original_size: Size of original content in bytes
            - text_size: Size of extracted text in characters
            - metadata: Additional extraction metadata (if include_metadata=True)

    Raises:
        ValueError: If URL is invalid or content type is not supported
        RuntimeError: If text extraction fails

    Example:
        >>> result = fetch_main_text("https://example.com/article.html")
        >>> print(result["text"])
        >>> print(f"Extracted {result['text_size']} characters from {result['content_type']}")

        >>> pdf_result = fetch_main_text("https://example.com/document.pdf")
        >>> print(pdf_result["text"])
    """
    # Fetch the URL content
    fetch_result = fetch_url(url, config, **kwargs)

    content_type = fetch_result["content_type"].lower()
    content = fetch_result["content"]

    # Get configuration options
    config = config or {}
    html_method = config.get("html_method", "readability")
    include_metadata = config.get("include_metadata", True)

    # Determine extraction method based on content type
    extraction_method = None
    text = None
    metadata = {}

    try:
        if content_type.startswith("text/html") or "html" in content_type:
            extraction_method = f"html_{html_method}"
            text = extract_html_text(content, html_method)
            if include_metadata:
                metadata["html_method"] = html_method

        elif content_type == "application/pdf" or content_type.endswith("/pdf"):
            extraction_method = "pdf"
            content_bytes = fetch_result["content_bytes"]
            text = extract_pdf_text(content_bytes)

        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or content_type.endswith("/docx")
        ):
            extraction_method = "docx"
            content_bytes = fetch_result["content_bytes"]
            text = extract_docx_text(content_bytes)

        elif content_type.startswith("text/"):
            extraction_method = "plain_text"
            text = content

        else:
            supported_types = [
                "text/html",
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "text/plain",
            ]
            raise ValueError(
                f"Unsupported content type: {content_type}. "
                f"Supported types: {', '.join(supported_types)}"
            )

    except Exception as e:
        if isinstance(e, (ValueError, RuntimeError)):
            raise
        raise RuntimeError(f"Text extraction failed: {e}")

    # Build result
    result: Dict[str, Any] = {
        "url": fetch_result["url"],
        "content_type": content_type,
        "text": text,
        "extraction_method": extraction_method,
        "original_size": fetch_result["size"],
        "text_size": len(text) if text else 0,
    }

    if include_metadata:
        result["metadata"] = {
            **metadata,
            "status_code": fetch_result["status_code"],
            "encoding": fetch_result["encoding"],
            "response_time": fetch_result["response_time"],
        }

    logger.info(
        f"Extracted {result['text_size']} characters using {extraction_method} from {content_type}"
    )
    return result
