"""Word document creation and modification functions for AI agents.

This module provides functions for creating and modifying Word (.docx) documents
using the python-docx library.
"""

import os

from ..decorators import strands_tool

try:
    from docx import Document  # type: ignore[import-untyped, import-not-found]

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


@strands_tool
def create_simple_docx(file_path: str, content: str, skip_confirm: bool) -> str:
    """Create simple Word document from text content.

    This function creates a basic Word document with the provided text content.
    The text will be formatted as a single paragraph.

    Args:
        file_path: Path where Word document will be created
        content: Text content to include in document
        skip_confirm: If False, raises error if file exists; if True, overwrites

    Returns:
        Success message with file path and size

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If file exists and skip_confirm=False, or content is empty
        PermissionError: If directory is not writable

    Example:
        >>> msg = create_simple_docx("/tmp/test.docx", "Hello World", False)
        >>> "Created Word document" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(content, str):
        raise TypeError("content must be a string")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    if not content.strip():
        raise ValueError("content cannot be empty")

    # Check if file exists
    if os.path.exists(file_path) and not skip_confirm:
        raise ValueError(
            f"File already exists: {file_path}. Set skip_confirm=True to overwrite."
        )

    # Check parent directory is writable
    parent_dir = os.path.dirname(file_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Parent directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        # Create document
        doc = Document()
        doc.add_paragraph(content)

        # Save document
        doc.save(file_path)

        # Get file size
        file_size = os.path.getsize(file_path)
        return f"Created Word document {file_path} ({file_size} bytes)"

    except Exception as e:
        raise ValueError(f"Failed to create Word document: {e}")


@strands_tool
def create_docx_from_paragraphs(
    file_path: str, paragraphs: list[str], skip_confirm: bool
) -> str:
    """Create Word document with multiple paragraphs.

    This function creates a Word document with multiple paragraphs.
    Each string in the list becomes a separate paragraph.

    Args:
        file_path: Path where Word document will be created
        paragraphs: List of text strings, one per paragraph
        skip_confirm: If False, raises error if file exists; if True, overwrites

    Returns:
        Success message with file path and paragraph count

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If file exists and skip_confirm=False, or paragraphs is empty

    Example:
        >>> paragraphs = ["Introduction", "Body text", "Conclusion"]
        >>> msg = create_docx_from_paragraphs("/tmp/test.docx", paragraphs, False)
        >>> "3 paragraphs" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(paragraphs, list):
        raise TypeError("paragraphs must be a list")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    if not paragraphs:
        raise ValueError("paragraphs cannot be empty")

    # Validate all paragraphs are strings
    for i, para in enumerate(paragraphs):
        if not isinstance(para, str):
            raise TypeError(f"Paragraph at index {i} must be a string")

    # Check if file exists
    if os.path.exists(file_path) and not skip_confirm:
        raise ValueError(
            f"File already exists: {file_path}. Set skip_confirm=True to overwrite."
        )

    # Check parent directory
    parent_dir = os.path.dirname(file_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Parent directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        doc = Document()

        # Add each paragraph
        for para_text in paragraphs:
            doc.add_paragraph(para_text)

        doc.save(file_path)

        return f"Created Word document {file_path} with {len(paragraphs)} paragraphs"

    except Exception as e:
        raise ValueError(f"Failed to create Word document: {e}")


@strands_tool
def create_docx_with_title(
    file_path: str, title: str, content: str, skip_confirm: bool
) -> str:
    """Create Word document with title and content.

    This function creates a Word document with a title using Heading 1 style
    followed by body content.

    Args:
        file_path: Path where Word document will be created
        title: Document title text (uses Heading 1 style)
        content: Body content text
        skip_confirm: If False, raises error if file exists; if True, overwrites

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If file exists and skip_confirm=False, or title/content empty

    Example:
        >>> msg = create_docx_with_title("/tmp/test.docx", "Report", "Content", False)
        >>> "Created Word document" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(title, str):
        raise TypeError("title must be a string")

    if not isinstance(content, str):
        raise TypeError("content must be a string")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    if not title.strip():
        raise ValueError("title cannot be empty")

    if not content.strip():
        raise ValueError("content cannot be empty")

    # Check if file exists
    if os.path.exists(file_path) and not skip_confirm:
        raise ValueError(
            f"File already exists: {file_path}. Set skip_confirm=True to overwrite."
        )

    # Check parent directory
    parent_dir = os.path.dirname(file_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Parent directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        doc = Document()

        # Add title with Heading 1 style
        doc.add_heading(title, level=1)

        # Add content
        doc.add_paragraph(content)

        doc.save(file_path)

        file_size = os.path.getsize(file_path)
        return f"Created Word document {file_path} ({file_size} bytes)"

    except Exception as e:
        raise ValueError(f"Failed to create Word document: {e}")


@strands_tool
def add_paragraph_to_docx(file_path: str, paragraph: str, skip_confirm: bool) -> str:
    """Append paragraph to existing Word document.

    This function adds a new paragraph to the end of an existing Word document,
    preserving all existing content.

    Args:
        file_path: Path to existing Word document
        paragraph: Text to add as new paragraph
        skip_confirm: If False, raises error if file doesn't exist; if True, creates new file

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If paragraph is empty
        FileNotFoundError: If file doesn't exist and skip_confirm=False

    Example:
        >>> msg = add_paragraph_to_docx("/tmp/test.docx", "New paragraph", True)
        >>> "Added paragraph" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(paragraph, str):
        raise TypeError("paragraph must be a string")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not paragraph.strip():
        raise ValueError("paragraph cannot be empty")

    # Check if file exists
    if not os.path.exists(file_path):
        if not skip_confirm:
            raise FileNotFoundError(
                f"File does not exist: {file_path}. Set skip_confirm=True to create new file."
            )
        # Create new document
        doc = Document()
    else:
        # Load existing document
        doc = Document(file_path)

    try:
        # Add paragraph
        doc.add_paragraph(paragraph)

        # Save document
        doc.save(file_path)

        return f"Added paragraph to {file_path}"

    except Exception as e:
        raise ValueError(f"Failed to add paragraph: {e}")


@strands_tool
def create_docx_with_headings(
    file_path: str, sections: list[dict[str, str]], skip_confirm: bool
) -> str:
    """Create Word document with headings and content sections.

    This function creates a structured Word document with headings and content.
    Each section can have a heading at a specific level and associated content.

    Args:
        file_path: Path where Word document will be created
        sections: List of section dictionaries, each with keys:
            - heading: Heading text (str)
            - level: Heading level 1-9 (str, e.g. "1", "2")
            - content: Body text for section (str)
        skip_confirm: If False, raises error if file exists; if True, overwrites

    Returns:
        Success message with file path and section count

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If sections invalid or file exists

    Example:
        >>> sections = [
        ...     {"heading": "Introduction", "level": "1", "content": "Intro text"},
        ...     {"heading": "Methods", "level": "1", "content": "Method details"}
        ... ]
        >>> msg = create_docx_with_headings("/tmp/test.docx", sections, False)
        >>> "2 sections" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(sections, list):
        raise TypeError("sections must be a list")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    if not sections:
        raise ValueError("sections cannot be empty")

    # Validate sections structure
    for i, section in enumerate(sections):
        if not isinstance(section, dict):
            raise TypeError(f"Section at index {i} must be a dictionary")
        if (
            "heading" not in section
            or "level" not in section
            or "content" not in section
        ):
            raise ValueError(
                f"Section at index {i} must have 'heading', 'level', and 'content' keys"
            )
        if not isinstance(section["heading"], str):
            raise TypeError(f"Section {i} heading must be a string")
        if not isinstance(section["level"], str):
            raise TypeError(f"Section {i} level must be a string")
        if not isinstance(section["content"], str):
            raise TypeError(f"Section {i} content must be a string")

        # Validate level is 1-9
        try:
            level_int = int(section["level"])
            if level_int < 1 or level_int > 9:
                raise ValueError(
                    f"Section {i} level must be between 1 and 9, got {level_int}"
                )
        except ValueError as e:
            raise ValueError(f"Section {i} level must be a valid integer: {e}")

    # Check if file exists
    if os.path.exists(file_path) and not skip_confirm:
        raise ValueError(
            f"File already exists: {file_path}. Set skip_confirm=True to overwrite."
        )

    # Check parent directory
    parent_dir = os.path.dirname(file_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Parent directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        doc = Document()

        # Add each section
        for section in sections:
            level = int(section["level"])
            doc.add_heading(section["heading"], level=level)
            if section["content"].strip():
                doc.add_paragraph(section["content"])

        doc.save(file_path)

        return f"Created Word document {file_path} with {len(sections)} sections"

    except Exception as e:
        raise ValueError(f"Failed to create Word document: {e}")


@strands_tool
def add_table_to_docx(
    file_path: str, table_data: list[list[str]], skip_confirm: bool
) -> str:
    """Append table to existing Word document.

    This function adds a table to the end of an existing Word document.
    The first row is typically used for headers.

    Args:
        file_path: Path to existing Word document
        table_data: 2D list representing table [row][column]
        skip_confirm: If False, raises error if file doesn't exist; if True, creates new file

    Returns:
        Success message with file path and table dimensions

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If table_data is empty or invalid
        FileNotFoundError: If file doesn't exist and skip_confirm=False

    Example:
        >>> table = [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
        >>> msg = add_table_to_docx("/tmp/test.docx", table, True)
        >>> "table" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(table_data, list):
        raise TypeError("table_data must be a list")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not table_data:
        raise ValueError("table_data cannot be empty")

    # Validate table structure
    for i, row in enumerate(table_data):
        if not isinstance(row, list):
            raise TypeError(f"Row at index {i} must be a list")
        for j, cell in enumerate(row):
            if not isinstance(cell, str):
                raise TypeError(f"Cell at row {i}, column {j} must be a string")

    # Check all rows have same length
    row_lengths = [len(row) for row in table_data]
    if len(set(row_lengths)) > 1:
        raise ValueError(f"All rows must have same length, got {row_lengths}")

    # Check if file exists
    if not os.path.exists(file_path):
        if not skip_confirm:
            raise FileNotFoundError(
                f"File does not exist: {file_path}. Set skip_confirm=True to create new file."
            )
        # Create new document
        doc = Document()
    else:
        # Load existing document
        doc = Document(file_path)

    try:
        # Add table
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)

        # Populate table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = cell_text

        # Save document
        doc.save(file_path)

        return f"Added {rows}x{cols} table to {file_path}"

    except Exception as e:
        raise ValueError(f"Failed to add table: {e}")


@strands_tool
def create_docx_from_template(
    template_path: str,
    output_path: str,
    replacements: dict[str, str],
    skip_confirm: bool,
) -> str:
    """Create Word document from template with placeholder replacements.

    This function loads a template document and replaces all instances of
    {{placeholder}} with corresponding values from the replacements dictionary.

    Args:
        template_path: Path to template Word document
        output_path: Path where filled document will be created
        replacements: Dictionary mapping placeholder names to replacement values
        skip_confirm: If False, raises error if output exists; if True, overwrites

    Returns:
        Success message with output path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If paths are invalid or output exists
        FileNotFoundError: If template doesn't exist

    Example:
        >>> replacements = {"name": "John", "date": "2024-01-15"}
        >>> msg = create_docx_from_template("/tmp/template.docx", "/tmp/output.docx", replacements, False)
        >>> "Created" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(template_path, str):
        raise TypeError("template_path must be a string")

    if not isinstance(output_path, str):
        raise TypeError("output_path must be a string")

    if not isinstance(replacements, dict):
        raise TypeError("replacements must be a dictionary")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not template_path.strip():
        raise ValueError("template_path cannot be empty")

    if not output_path.strip():
        raise ValueError("output_path cannot be empty")

    # Validate replacements
    for key, value in replacements.items():
        if not isinstance(key, str):
            raise TypeError(f"Replacement key must be string, got {type(key)}")
        if not isinstance(value, str):
            raise TypeError(f"Replacement value must be string, got {type(value)}")

    # Check template exists
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Check if output exists
    if os.path.exists(output_path) and not skip_confirm:
        raise ValueError(
            f"Output file already exists: {output_path}. Set skip_confirm=True to overwrite."
        )

    # Check output directory
    parent_dir = os.path.dirname(output_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Output directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        # Load template
        doc = Document(template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # Replace in full paragraph text
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)

        # Save filled document
        doc.save(output_path)

        file_size = os.path.getsize(output_path)
        return f"Created document from template: {output_path} ({file_size} bytes)"

    except Exception as e:
        raise ValueError(f"Failed to create document from template: {e}")


@strands_tool
def docx_to_text(file_path: str, output_path: str, skip_confirm: bool) -> str:
    """Convert Word document to plain text file.

    This function extracts all text from a Word document and saves it
    to a plain text file, stripping all formatting.

    Args:
        file_path: Path to Word document to convert
        output_path: Path where text file will be created
        skip_confirm: If False, raises error if output exists; if True, overwrites

    Returns:
        Success message with output path and size

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If paths are invalid or output exists
        FileNotFoundError: If input doesn't exist

    Example:
        >>> msg = docx_to_text("/tmp/document.docx", "/tmp/output.txt", False)
        >>> "Converted" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(output_path, str):
        raise TypeError("output_path must be a string")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    if not output_path.strip():
        raise ValueError("output_path cannot be empty")

    # Check input exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Input document not found: {file_path}")

    # Check if output exists
    if os.path.exists(output_path) and not skip_confirm:
        raise ValueError(
            f"Output file already exists: {output_path}. Set skip_confirm=True to overwrite."
        )

    # Check output directory
    parent_dir = os.path.dirname(output_path) or "."
    if not os.path.exists(parent_dir):
        raise FileNotFoundError(f"Output directory does not exist: {parent_dir}")

    if not os.access(parent_dir, os.W_OK):
        raise PermissionError(f"Cannot write to directory: {parent_dir}")

    try:
        # Load document
        doc = Document(file_path)

        # Extract all text
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Write to text file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(text_parts))

        file_size = os.path.getsize(output_path)
        return f"Converted {file_path} to text file: {output_path} ({file_size} bytes)"

    except Exception as e:
        raise ValueError(f"Failed to convert document to text: {e}")
