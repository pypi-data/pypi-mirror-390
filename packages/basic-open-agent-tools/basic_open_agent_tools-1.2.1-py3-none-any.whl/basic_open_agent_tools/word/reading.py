"""Word document reading and extraction functions for AI agents.

This module provides functions for reading Word (.docx) files, extracting text,
tables, metadata, and searching within documents.
"""

import os

from ..decorators import strands_tool

try:
    from docx import Document  # type: ignore[import-untyped, import-not-found]

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


# Maximum file size to prevent memory exhaustion (50MB default)
MAX_DOCX_FILE_SIZE = 50 * 1024 * 1024


@strands_tool
def extract_text_from_docx(file_path: str) -> str:
    """Extract all text content from Word document.

    This function reads a Word document and extracts all text from every
    paragraph, concatenating it into a single string. Requires python-docx.

    Args:
        file_path: Path to Word document to read

    Returns:
        Extracted text content from all paragraphs concatenated together

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is too large or not a valid Word document
        PermissionError: If file cannot be read
        TypeError: If file_path is not a string

    Example:
        >>> text = extract_text_from_docx("/data/document.docx")
        >>> "Introduction" in text
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    # Check file size
    file_size = os.path.getsize(file_path)
    if file_size > MAX_DOCX_FILE_SIZE:
        raise ValueError(
            f"Word document too large: {file_size} bytes "
            f"(maximum: {MAX_DOCX_FILE_SIZE} bytes)"
        )

    # Check read permission
    if not os.access(file_path, os.R_OK):
        raise PermissionError(f"Cannot read Word document: {file_path}")

    try:
        doc = Document(file_path)

        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        return "\n".join(text_parts)

    except Exception as e:
        raise ValueError(f"Failed to read Word document {file_path}: {e}")


@strands_tool
def get_docx_paragraphs(file_path: str) -> list[str]:
    """Get all paragraphs from Word document as separate strings.

    This function extracts each paragraph as a separate string,
    preserving the document's paragraph structure.

    Args:
        file_path: Path to Word document to read

    Returns:
        List of paragraph text strings

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is malformed
        TypeError: If file_path is not a string

    Example:
        >>> paragraphs = get_docx_paragraphs("/data/document.docx")
        >>> len(paragraphs) > 0
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    # Check file size
    file_size = os.path.getsize(file_path)
    if file_size > MAX_DOCX_FILE_SIZE:
        raise ValueError(
            f"Word document too large: {file_size} bytes "
            f"(maximum: {MAX_DOCX_FILE_SIZE} bytes)"
        )

    try:
        doc = Document(file_path)

        # Extract all paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            paragraphs.append(paragraph.text)

        return paragraphs

    except Exception as e:
        raise ValueError(f"Failed to read Word document {file_path}: {e}")


@strands_tool
def get_docx_tables(file_path: str) -> list[list[list[str]]]:
    """Extract all tables from Word document.

    This function extracts all tables from the document as a 3D structure:
    [table_index][row_index][cell_index].

    Args:
        file_path: Path to Word document to read

    Returns:
        List of tables, each table is a list of rows, each row is a list of cell values

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is malformed
        TypeError: If file_path is not a string

    Example:
        >>> tables = get_docx_tables("/data/document.docx")
        >>> len(tables) >= 0
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        doc = Document(file_path)

        # Extract all tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)

        return tables

    except Exception as e:
        raise ValueError(f"Failed to extract tables from {file_path}: {e}")


@strands_tool
def get_docx_metadata(file_path: str) -> dict[str, str]:
    """Extract metadata from Word document.

    This function reads Word document metadata including author, title,
    subject, keywords, created date, and modified date.

    Args:
        file_path: Path to Word document to read

    Returns:
        Dictionary with metadata fields (all values are strings):
        - author: Document author
        - title: Document title
        - subject: Document subject
        - keywords: Document keywords
        - created: Creation date/time
        - modified: Last modification date/time

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is malformed
        TypeError: If file_path is not a string

    Example:
        >>> metadata = get_docx_metadata("/data/document.docx")
        >>> "author" in metadata
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        doc = Document(file_path)
        core_props = doc.core_properties

        result: dict[str, str] = {
            "author": str(core_props.author) if core_props.author else "",
            "title": str(core_props.title) if core_props.title else "",
            "subject": str(core_props.subject) if core_props.subject else "",
            "keywords": str(core_props.keywords) if core_props.keywords else "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        return result

    except Exception as e:
        raise ValueError(f"Failed to extract metadata from {file_path}: {e}")


@strands_tool
def search_docx_text(
    file_path: str, search_term: str, case_sensitive: bool
) -> list[dict[str, object]]:
    """Search for text in Word document and return matches.

    This function searches all paragraphs for the specified text and returns
    matches with context. Each match includes paragraph index and text.

    Args:
        file_path: Path to Word document to search
        search_term: Text to search for
        case_sensitive: Whether search should be case-sensitive

    Returns:
        List of match dictionaries, each containing:
        - paragraph_index: Index of paragraph where match was found (int, 0-indexed)
        - match_text: The matched text (str)
        - paragraph_text: Full paragraph text containing match (str)

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If search_term is empty or file is malformed
        TypeError: If parameters are wrong type

    Example:
        >>> matches = search_docx_text("/data/document.docx", "Python", False)
        >>> len(matches) >= 0
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(search_term, str):
        raise TypeError("search_term must be a string")

    if not isinstance(case_sensitive, bool):
        raise TypeError("case_sensitive must be a boolean")

    if not search_term.strip():
        raise ValueError("search_term cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        doc = Document(file_path)
        matches: list[dict[str, object]] = []

        for para_index, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text

            # Prepare text for searching
            search_text = para_text if case_sensitive else para_text.lower()
            search_for = search_term if case_sensitive else search_term.lower()

            if search_for in search_text:
                matches.append(
                    {
                        "paragraph_index": para_index,
                        "match_text": search_term,
                        "paragraph_text": para_text,
                    }
                )

        return matches

    except Exception as e:
        raise ValueError(f"Failed to search Word document {file_path}: {e}")


@strands_tool
def get_docx_info(file_path: str) -> dict[str, object]:
    """Get comprehensive information about Word document.

    This function returns detailed information about the Word document including
    paragraph count, table count, file size, and metadata.

    Args:
        file_path: Path to Word document to analyze

    Returns:
        Dictionary with document information:
        - paragraph_count: Number of paragraphs (int)
        - table_count: Number of tables (int)
        - file_size_bytes: File size in bytes (int)
        - metadata: Dictionary of metadata fields (dict[str, str])

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is malformed
        TypeError: If file_path is not a string

    Example:
        >>> info = get_docx_info("/data/document.docx")
        >>> info["paragraph_count"] >= 0
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not file_path.strip():
        raise ValueError("file_path cannot be empty")

    # Validate file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        doc = Document(file_path)
        file_size = os.path.getsize(file_path)

        # Get metadata
        core_props = doc.core_properties
        metadata_dict: dict[str, str] = {
            "author": str(core_props.author) if core_props.author else "",
            "title": str(core_props.title) if core_props.title else "",
            "subject": str(core_props.subject) if core_props.subject else "",
            "keywords": str(core_props.keywords) if core_props.keywords else "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        return {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size_bytes": file_size,
            "metadata": metadata_dict,
        }

    except Exception as e:
        raise ValueError(f"Failed to get document info from {file_path}: {e}")
