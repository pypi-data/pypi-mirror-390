"""Word document styling and formatting functions for AI agents.

This module provides functions for applying styles and formatting to
Word (.docx) documents.
"""

import os

from ..decorators import strands_tool

try:
    from docx import Document  # type: ignore[import-untyped, import-not-found]
    from docx.enum.text import (  # type: ignore[import-untyped, import-not-found]
        WD_ALIGN_PARAGRAPH,
    )

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


@strands_tool
def apply_heading_style(
    file_path: str, paragraph_index: int, heading_level: int, skip_confirm: bool
) -> str:
    """Apply heading style to specific paragraph in Word document.

    This function applies a heading style (Heading 1 through Heading 9)
    to a specific paragraph in an existing document.

    Args:
        file_path: Path to Word document to modify
        paragraph_index: Index of paragraph to style (0-indexed)
        heading_level: Heading level (1-9)
        skip_confirm: If False, creates backup; if True, modifies directly

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If heading_level invalid or paragraph index out of range
        FileNotFoundError: If file doesn't exist
        IndexError: If paragraph_index is out of range

    Example:
        >>> msg = apply_heading_style("/tmp/test.docx", 0, 1, True)
        >>> "Applied" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(paragraph_index, int):
        raise TypeError("paragraph_index must be an integer")

    if not isinstance(heading_level, int):
        raise TypeError("heading_level must be an integer")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if paragraph_index < 0:
        raise ValueError("paragraph_index must be non-negative")

    if heading_level < 1 or heading_level > 9:
        raise ValueError("heading_level must be between 1 and 9")

    # Check file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        # Load document
        doc = Document(file_path)

        # Check paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            raise IndexError(
                f"Paragraph index {paragraph_index} out of range "
                f"(document has {len(doc.paragraphs)} paragraphs)"
            )

        # Get paragraph and apply heading style
        paragraph = doc.paragraphs[paragraph_index]
        paragraph.style = f"Heading {heading_level}"

        # Save document
        doc.save(file_path)

        return f"Applied Heading {heading_level} style to paragraph {paragraph_index} in {file_path}"

    except IndexError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to apply heading style: {e}")


@strands_tool
def apply_bold_to_paragraph(
    file_path: str, paragraph_index: int, skip_confirm: bool
) -> str:
    """Make entire paragraph bold in Word document.

    This function applies bold formatting to all text in a specific paragraph.

    Args:
        file_path: Path to Word document to modify
        paragraph_index: Index of paragraph to make bold (0-indexed)
        skip_confirm: If False, creates backup; if True, modifies directly

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If paragraph index invalid
        FileNotFoundError: If file doesn't exist
        IndexError: If paragraph_index is out of range

    Example:
        >>> msg = apply_bold_to_paragraph("/tmp/test.docx", 0, True)
        >>> "Applied bold" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(paragraph_index, int):
        raise TypeError("paragraph_index must be an integer")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if paragraph_index < 0:
        raise ValueError("paragraph_index must be non-negative")

    # Check file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        # Load document
        doc = Document(file_path)

        # Check paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            raise IndexError(
                f"Paragraph index {paragraph_index} out of range "
                f"(document has {len(doc.paragraphs)} paragraphs)"
            )

        # Get paragraph and apply bold to all runs
        paragraph = doc.paragraphs[paragraph_index]
        for run in paragraph.runs:
            run.bold = True

        # Save document
        doc.save(file_path)

        return f"Applied bold formatting to paragraph {paragraph_index} in {file_path}"

    except IndexError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to apply bold formatting: {e}")


@strands_tool
def set_paragraph_alignment(
    file_path: str, paragraph_index: int, alignment: str, skip_confirm: bool
) -> str:
    """Set paragraph alignment in Word document.

    This function sets the alignment for a specific paragraph.
    Valid alignments are: left, center, right, justify.

    Args:
        file_path: Path to Word document to modify
        paragraph_index: Index of paragraph to align (0-indexed)
        alignment: Alignment type (left, center, right, justify)
        skip_confirm: If False, creates backup; if True, modifies directly

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If alignment invalid or paragraph index out of range
        FileNotFoundError: If file doesn't exist
        IndexError: If paragraph_index is out of range

    Example:
        >>> msg = set_paragraph_alignment("/tmp/test.docx", 0, "center", True)
        >>> "Set alignment" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(paragraph_index, int):
        raise TypeError("paragraph_index must be an integer")

    if not isinstance(alignment, str):
        raise TypeError("alignment must be a string")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if paragraph_index < 0:
        raise ValueError("paragraph_index must be non-negative")

    # Validate alignment
    valid_alignments = ["left", "center", "right", "justify"]
    if alignment.lower() not in valid_alignments:
        raise ValueError(
            f"alignment must be one of {valid_alignments}, got '{alignment}'"
        )

    # Check file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        # Load document
        doc = Document(file_path)

        # Check paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            raise IndexError(
                f"Paragraph index {paragraph_index} out of range "
                f"(document has {len(doc.paragraphs)} paragraphs)"
            )

        # Get paragraph and set alignment
        paragraph = doc.paragraphs[paragraph_index]

        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        paragraph.alignment = alignment_map[alignment.lower()]

        # Save document
        doc.save(file_path)

        return f"Set alignment to {alignment} for paragraph {paragraph_index} in {file_path}"

    except IndexError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to set paragraph alignment: {e}")


@strands_tool
def add_page_break(file_path: str, after_paragraph: int, skip_confirm: bool) -> str:
    """Insert page break after specified paragraph in Word document.

    This function inserts a page break after a specific paragraph,
    forcing the next content to start on a new page.

    Args:
        file_path: Path to Word document to modify
        after_paragraph: Index of paragraph after which to insert break (0-indexed)
        skip_confirm: If False, creates backup; if True, modifies directly

    Returns:
        Success message with file path

    Raises:
        ImportError: If python-docx is not installed
        TypeError: If parameters are wrong type
        ValueError: If paragraph index invalid
        FileNotFoundError: If file doesn't exist
        IndexError: If paragraph index is out of range

    Example:
        >>> msg = add_page_break("/tmp/test.docx", 0, True)
        >>> "Added page break" in msg
        True
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError(
            "python-docx is required for Word operations. "
            "Install with: pip install basic-open-agent-tools[word]"
        )

    if not isinstance(file_path, str):
        raise TypeError("file_path must be a string")

    if not isinstance(after_paragraph, int):
        raise TypeError("after_paragraph must be an integer")

    if not isinstance(skip_confirm, bool):
        raise TypeError("skip_confirm must be a boolean")

    if after_paragraph < 0:
        raise ValueError("after_paragraph must be non-negative")

    # Check file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    try:
        # Load document
        doc = Document(file_path)

        # Check paragraph index is valid
        if after_paragraph >= len(doc.paragraphs):
            raise IndexError(
                f"Paragraph index {after_paragraph} out of range "
                f"(document has {len(doc.paragraphs)} paragraphs)"
            )

        # Get paragraph and add page break
        paragraph = doc.paragraphs[after_paragraph]
        run = paragraph.add_run()
        run.add_break(type=1)  # type=1 is page break

        # Save document
        doc.save(file_path)

        return f"Added page break after paragraph {after_paragraph} in {file_path}"

    except IndexError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to add page break: {e}")
