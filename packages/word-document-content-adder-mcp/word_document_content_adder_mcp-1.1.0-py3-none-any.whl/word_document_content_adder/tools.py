"""
Word文档内容添加工具函数

提供五个核心的文档内容添加功能
"""

import os
from typing import Optional, List
from .utils import (
    ensure_docx_extension,
    check_file_writeable,
    get_default_desktop_dir
)


async def create_document(output_filename: str, template: Optional[str] = None, overwrite: bool = False, output_dir: Optional[str] = None) -> str:
    """Create a new Word (.docx) document at a specified output directory (default Desktop).

    Args:
        output_filename: 输出文件名（仅文件名，自动补全 .docx 扩展名，忽略路径部分）
        template: 模板文档路径（可选，基于此模板创建新文档）
        overwrite: 是否允许覆盖已存在文件（默认 False）
        output_dir: 输出目录路径（可选，默认为用户桌面，自动识别 OneDrive 或本地桌面）
    """
    from docx import Document

    # Ensure .docx extension
    output_filename = ensure_docx_extension(output_filename)

    # Determine output directory (default to Desktop via resolver)
    try:
        base_dir = output_dir.strip() if output_dir else ""
    except Exception:
        base_dir = ""
    if not base_dir:
        base_dir = get_default_desktop_dir()

    # Enforce filename as basename (no directories)
    output_filename = os.path.basename(output_filename)

    # Build target path: always join base_dir with output_filename
    target_path = os.path.join(base_dir, output_filename)

    # Ensure target directory exists
    try:
        dir_name = os.path.dirname(target_path)
        if dir_name:
            os.makedirs(dir_name, exist_ok=True)
    except Exception:
        # Non-fatal: directory creation failure will be surfaced by save()
        pass

    # Handle overwrite / existence
    if os.path.exists(target_path):
        # 若文件已存在且未要求覆盖，则改为幂等行为：不报错，直接返回提示信息
        if not overwrite:
            abs_target = os.path.abspath(target_path)
            return (
                "检测到现有文档，未创建新文件。\n"
                f"输出路径: {abs_target}\n\n"
                "请在后续工具调用中使用完整路径以添加内容: "
                f"{abs_target}"
            )
        # 需要覆盖时检查可写性
        is_writeable, error_message = check_file_writeable(target_path)
        if not is_writeable:
            raise PermissionError(f"Cannot overwrite document: {error_message}. Close the file or choose a different path.")

    # Create document (from template or blank)
    try:
        if template:
            abs_template = os.path.abspath(template)
            if not os.path.exists(abs_template):
                raise FileNotFoundError(f"Template file not found: {abs_template}")
            try:
                doc = Document(abs_template)
            except Exception as open_error:
                try:
                    from docx.opc.exceptions import PackageNotFoundError
                except Exception:
                    PackageNotFoundError = tuple()  # type: ignore
                if isinstance(open_error, PackageNotFoundError):
                    raise RuntimeError(
                        f"Template is not a valid .docx package or is corrupted: {abs_template}."
                    )
                raise
        else:
            doc = Document()

        # Save to target path
        abs_target = os.path.abspath(target_path)
        doc.save(abs_target)
        
        # Return message emphasizing the output path for subsequent tool calls
        template_info = f" using template {os.path.abspath(template)}" if template else ""
        return f"文档创建成功！\n输出路径: {abs_target}\n\n请在后续工具调用中使用完整路径: {abs_target}{template_info}"
    except Exception as e:
        raise RuntimeError(f"Failed to create document: {str(e)}")

async def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    filename = ensure_docx_extension(filename)
    
    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: level must be an integer between 1 and 9")
    
    # Validate level range
    if level < 1 or level > 9:
        raise ValueError(f"Invalid heading level: {level}. Level must be between 1 and 9.")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        # 捕获无效docx包的错误，提供更清晰提示
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        # 统一使用段落样式方式，并在缺少样式时创建；先尝试升级末段，避免重复插入
        style_name = 'Heading ' + str(level)
        try:
            _ = doc.styles[style_name]
        except KeyError:
            # 创建基本标题样式（英文名称），以适配不同本地化模板
            new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            try:
                new_style.base_style = doc.styles['Normal']
            except Exception:
                pass
            new_style.font.bold = True
            if level == 1:
                new_style.font.size = Pt(16)
            elif level == 2:
                new_style.font.size = Pt(14)
            else:
                new_style.font.size = Pt(12)

        # 如果文档末段文本与将插入的标题相同，则直接将该段落升级为标题，避免“正文+标题”重复
        if doc.paragraphs and doc.paragraphs[-1].text.strip() == str(text).strip():
            paragraph = doc.paragraphs[-1]
            try:
                paragraph.style = style_name
                doc.save(filename)
                return f"Upgraded last paragraph to heading '{text}' (level {level}) in {filename}"
            except Exception:
                # 若设置样式失败，则进行直接格式化作为兜底
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(16)
                    elif level == 2:
                        run.font.size = Pt(14)
                    else:
                        run.font.size = Pt(12)
                doc.save(filename)
                return f"Upgraded last paragraph with direct formatting for heading '{text}' in {filename}"

        # 正常路径：插入一个标题段落（使用样式）
        paragraph = doc.add_paragraph(text)
        try:
            paragraph.style = style_name
        except Exception:
            # 样式设置失败则直接格式化
            if paragraph.runs:
                run = paragraph.runs[0]
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add heading: {str(e)}")


async def add_paragraph(filename: str, text: str, style: Optional[str] = None) -> str:
    """Add a paragraph to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        paragraph = doc.add_paragraph(text)
        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add paragraph: {str(e)}")


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        table = doc.add_table(rows=rows, cols=cols)
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        raise RuntimeError(f"Failed to add table: {str(e)}")


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.

    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)

    行为说明：
    - 当未提供 `width` 时，默认使用图片原始尺寸；
    - 若图片原始宽度超过文档内容区域（页面宽度扣除左右页边距），将自动按文档内容宽度 100% 进行缩放；
    - 当提供 `width` 时，按指定宽度（英寸）比例缩放。
    """
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    filename = ensure_docx_extension(filename)

    # Validate document existence
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)

    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        raise FileNotFoundError(f"Image file not found: {abs_image_path}")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(abs_filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {abs_filename}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        diagnostic = f"Document: {abs_filename}, Image: {abs_image_path}"
        try:
            # 计算文档内容区域宽度（英寸）
            try:
                section = doc.sections[0] if doc.sections else None
                content_width_inches = None
                if section is not None:
                    content_width_inches = (
                        section.page_width.inches - (section.left_margin.inches + section.right_margin.inches)
                    )
            except Exception:
                content_width_inches = None

            # 计算图片原始宽度（英寸），若无法获取 DPI 则回退到 96dpi
            try:
                with Image.open(abs_image_path) as im:
                    width_px = im.width
                    dpi_info = im.info.get("dpi", (96, 96))
                    dpi_x = dpi_info[0] if isinstance(dpi_info, tuple) and len(dpi_info) >= 1 else 96
                    if not dpi_x or dpi_x <= 1:
                        dpi_x = 96
                    image_width_inches = float(width_px) / float(dpi_x)
            except Exception:
                image_width_inches = None

            # 决定插入时使用的宽度
            width_arg = None
            if width is not None and width > 0:
                width_arg = Inches(width)
            else:
                # 未指定宽度时，若图片过宽则缩放到内容区域宽度
                if (
                    content_width_inches is not None
                    and image_width_inches is not None
                    and image_width_inches > content_width_inches
                ):
                    width_arg = Inches(content_width_inches)

            if width_arg is not None:
                doc.add_picture(abs_image_path, width=width_arg)
            else:
                doc.add_picture(abs_image_path)

            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            raise RuntimeError(f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}")
    except Exception as outer_error:
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        raise RuntimeError(f"Document processing error: {error_type} - {error_msg or 'No error details available'}")


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.

    Args:
        filename: Path to the Word document
    """
    from docx import Document

    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        raise RuntimeError(f"Failed to add page break: {str(e)}")
