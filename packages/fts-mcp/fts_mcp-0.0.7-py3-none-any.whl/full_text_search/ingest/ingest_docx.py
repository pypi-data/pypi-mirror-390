#!/usr/bin/env python3
"""
Convert an code DOCX to JSONL, one record per Word section.
Each JSON object has a single key, "text", containing the section’s full text
(including numbering/lettering).
"""

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentCls
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .common import coalesce_small_chunks


# ----------------------------- helpers --------------------------------- #
def iter_block_items(doc: DocumentCls):
    """
    Yield Paragraph or Table objects in document order.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def table_to_text(tbl, sep: str = "\t") -> str:
    """
    Robustly flatten a Word table to text, tolerating vertical merges.
    Each row becomes one line; cells joined with `sep`.
    """
    lines: list[str] = []

    # tbl._tbl.tr_lst is the list of <w:tr> elements in order
    for tr in tbl._tbl.tr_lst:
        # tr.tc_lst is every <w:tc> actually present in this row
        cell_texts = []
        for tc in tr.tc_lst:
            cell_texts.append(_Cell(tc, tbl).text.strip())
        lines.append(sep.join(cell_texts))

    return "\n".join(lines)


def roman(n: int) -> str:
    vals = zip(
        (1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1),
        ("M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"),
    )
    out = []
    for v, s in vals:
        q, n = divmod(n, v)
        out.append(s * q)
    return "".join(out)


def next_marker(fmt: str, counter: int) -> str:
    if fmt in ("decimal", "decimalZero"):
        return str(counter)
    if fmt == "lowerLetter":
        return chr(ord("a") + counter - 1)
    if fmt == "upperLetter":
        return chr(ord("A") + counter - 1)
    if fmt == "lowerRoman":
        return roman(counter).lower()
    if fmt == "upperRoman":
        return roman(counter)
    return str(counter)  # fallback


# --------------------- numbering definition table ---------------------- #
def build_num_table(doc: DocumentCls):
    """
    Return dict[(numId, ilvl)] -> (fmt, lvlText, start)
    """
    tbl = {}
    num_root = doc.part.numbering_part.element  # type: ignore
    for num in num_root.findall(".//w:num", num_root.nsmap):
        numId = num.get(qn("w:numId"))
        abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
        abs_xpath = f'.//w:abstractNum[@w:abstractNumId="{abstract_id}"]'
        abstract = num_root.find(abs_xpath, num_root.nsmap)

        for lvl in abstract.findall("w:lvl", num_root.nsmap):
            ilvl = lvl.get(qn("w:ilvl"))
            fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
            text = lvl.find(qn("w:lvlText")).get(qn("w:val"))
            start_el = lvl.find(qn("w:start"))
            start = int(start_el.get(qn("w:val"))) if start_el is not None else 1
            tbl[(numId, ilvl)] = (fmt, text, start)
    return tbl


# --------------------------- main iterator ----------------------------- #
def paragraphs_with_prefix(doc: DocumentCls):
    """
    Yield (prefix + text, paragraph) for every paragraph in document,
    where prefix is the list or outline marker (if any).
    """
    num_tbl = build_num_table(doc)
    counters = {}  # (numId, ilvl) -> current int

    for p in doc.paragraphs:
        pr = p._p.pPr
        prefix = ""
        if pr is not None and pr.numPr is not None:
            numId = pr.numPr.numId.val
            ilvl = pr.numPr.ilvl.val
            key = (numId, ilvl)
            fmt, lvl_text, start = num_tbl.get(key, ("decimal", "%1.", 1))
            count = counters.get(key, start)
            marker = next_marker(fmt, count)
            prefix = re.sub(r"%\d", marker, lvl_text).strip() + " "
            counters[key] = count + 1
        yield prefix + p.text, p


def blocks_with_prefix(doc: DocumentCls):
    """
    Yield (text, is_section_break) for every *block* (paragraph or table),
    inserting list/outline markers in front of paragraph text.
    Inline tables are flattened via `table_to_text`.
    """
    num_tbl = build_num_table(doc)
    counters: dict[tuple[str, str], int] = {}

    for blk in iter_block_items(doc):
        if isinstance(blk, Paragraph):
            pr = blk._p.pPr
            prefix = ""
            if pr is not None and pr.numPr is not None:
                key = (pr.numPr.numId.val, pr.numPr.ilvl.val)
                fmt, lvl_text, start = num_tbl.get(key, ("decimal", "%1.", 1))
                count = counters.get(key, start)
                marker = next_marker(fmt, count)  # type: ignore
                prefix = re.sub(r"%\d", marker, lvl_text).strip() + " "
                counters[key] = count + 1  # type: ignore

            text = prefix + blk.text.strip()
            sect_break = pr is not None and pr.sectPr is not None
            yield text, sect_break

        else:  # Table
            yield table_to_text(blk), False


def paragraphs(doc: DocumentCls):
    """
    Just yields paragraph for every paragraph in the document,
    for codes that are not numbered with Word numbering.
    """
    for p in doc.paragraphs:
        yield p.text.strip()


# ---------------------------- conversion -------------------------------- #
def docx_to_jsonl_sections(in_path: Path | str):
    """
    For codes like Oakland (Municode) with section breaks
    separating logical sections.
    """
    doc = Document(str(in_path))
    buffer, sections = [], []

    for text, p in paragraphs_with_prefix(doc):
        buffer.append(text)
        sect_break = p._p.pPr is not None and p._p.pPr.sectPr is not None
        if sect_break:  # end-of-section
            sections.append("\n".join(buffer).strip())
            buffer.clear()

    if buffer:  # trailing material
        sections.append("\n".join(buffer).strip())

    sections = coalesce_small_chunks(sections)

    return sections


def docx_to_jsonl_sections_with_tables(in_path: Path | str):
    doc = Document(str(in_path))
    buffer, sections = [], []

    for text, sect_break in blocks_with_prefix(doc):
        buffer.append(text)
        if sect_break:
            sections.append("\n".join(buffer).strip())
            buffer.clear()

    if buffer:  # trailing material
        sections.append("\n".join(buffer).strip())

    sections = coalesce_small_chunks(sections)

    return sections


def docx_to_jsonl_lines(in_path: Path | str, split_on: list[str]):
    """
    For codes like Vail (AmLegal) where everything is paragraphs,
    must provide prefixes which signal the start of a new section.
    """
    doc = Document(str(in_path))
    buffer, sections = [], []

    for text in paragraphs(doc):
        if any(text.strip().startswith(x) for x in split_on):
            if buffer:
                sections.append("\n\n".join(buffer).strip())
                buffer.clear()

        buffer.append(text)

    if buffer:  # trailing material
        sections.append("\n".join(buffer).strip())

    sections = coalesce_small_chunks(sections)
    return sections
