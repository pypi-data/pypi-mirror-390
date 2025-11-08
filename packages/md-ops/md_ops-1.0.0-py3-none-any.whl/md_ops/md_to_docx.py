"""
Convert Markdown to DOCX
"""

import markdown2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from html.parser import HTMLParser


class MarkdownToDocxConverter(HTMLParser):
    def __init__(self, document):
        super().__init__()
        self.doc = document
        self.current_paragraph = None
        self.current_run = None
        self.list_level = 0
        self.in_code_block = False
        self.in_blockquote = False
        self.heading_level = 0
        
    def handle_starttag(self, tag, attrs):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_heading('', level=self.heading_level)
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'strong' or tag == 'b':
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                self.current_run.bold = True
        elif tag == 'em' or tag == 'i':
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                self.current_run.italic = True
        elif tag == 'code':
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                self.current_run.font.name = 'Courier New'
        elif tag == 'pre':
            self.in_code_block = True
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = 'Intense Quote'
        elif tag == 'blockquote':
            self.in_blockquote = True
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = 'Quote'
        elif tag == 'ul' or tag == 'ol':
            self.list_level += 1
        elif tag == 'li':
            self.current_paragraph = self.doc.add_paragraph(style='List Bullet' if self.list_level > 0 else 'Normal')
        elif tag == 'hr':
            self.doc.add_paragraph('_' * 50)
        elif tag == 'a':
            # Handle links
            href = dict(attrs).get('href', '')
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                
    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = 0
        elif tag in ['strong', 'b', 'em', 'i', 'code']:
            self.current_run = None
        elif tag == 'pre':
            self.in_code_block = False
        elif tag == 'blockquote':
            self.in_blockquote = False
        elif tag in ['ul', 'ol']:
            self.list_level -= 1
            
    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.current_paragraph:
                if self.current_run:
                    self.current_run.text = data
                else:
                    self.current_paragraph.add_run(data)


def md_to_docx(markdown_text, output_path="output.docx"):
    """
    Convert Markdown text to a DOCX file.
    
    Args:
        markdown_text (str): The Markdown content to convert
        output_path (str): Output file path (default: 'output.docx')
        
    Returns:
        str: Path to the created DOCX file
    """
    # Convert markdown to HTML
    html = markdown2.markdown(markdown_text, extras=['fenced-code-blocks', 'tables', 'break-on-newline'])
    
    # Create a new Document
    doc = Document()
    
    # Parse HTML and convert to DOCX
    converter = MarkdownToDocxConverter(doc)
    converter.feed(html)
    
    # Save the document
    doc.save(output_path)
    
    return output_path
