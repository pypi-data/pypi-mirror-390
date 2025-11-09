# -*- coding: utf-8 -*-
"""
本模块功能：转换ipynb文件为docx，带有的目录，代码行加边框，图像适配页宽。
注意：
    需要安装pandoc并将其路径加入操作系统的PATH。
    可在Anaconda Prompt或macOS Terminal下输入pandoc尝试，若未加入PATH则提示找不到。
尚存问题：
    1. 标题行未居中，且重复生成；
    2. 目录页码不准确，需要手动更新；
    3. 若docx文件已打开出错。
所属工具包：证券投资分析工具SIAT 
SIAT：Security Investment Analysis Tool
创建日期：2025年7月8日
最新修订日期：2025年7月8日
作者：王德宏 (WANG Dehong, Peter)
作者单位：北京外国语大学国际商学院
作者邮件：wdehong2000@163.com
版权所有：王德宏
用途限制：仅限研究与教学使用。
特别声明：作者不对使用本工具进行证券投资导致的任何损益负责！
"""

#==============================================================================
#关闭所有警告
import warnings; warnings.filterwarnings('ignore')

import os,sys
import errno
import tempfile
import subprocess

import nbformat
from nbconvert import HTMLExporter
import pypandoc

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph

import logging

import contextlib
import io

import time
from IPython.display import display, Javascript

# —— 新增：Notebook 强制保存：貌似不管用 ——
def _save_current_notebook():
    """
    在浏览器端触发一次保存：兼容 Classic Notebook、Lab 3.x/4.x。
    """
    js = """
    (function() {
      // Classic Notebook
      if (window.Jupyter && Jupyter.notebook) {
        Jupyter.notebook.save_checkpoint();
      }
      // JupyterLab >=3: 用 app.commands
      else if (window.jupyterapp && jupyterapp.commands) {
        jupyterapp.commands.execute('docmanager:save');
      }
      // JupyterLab <=2 或其他
      else if (window.require) {
        require(['@jupyterlab/docmanager'], function(docManager) {
          docManager.save();
        });
      }
    })();
    """
    try:
        display(Javascript(js))
        time.sleep(0.5)   # 给浏览器一点时间写盘
    except Exception:
        pass


# 预设纸张尺寸（单位：毫米）
PAGE_SIZES = {"A4": (210, 297), "A3": (297, 420)}


def _add_border_to_paragraph(paragraph):
    """给 paragraph 添加四边单线边框"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        elm = OxmlElement(f'w:{edge}')
        elm.set(qn('w:val'), 'single')
        elm.set(qn('w:sz'), '4')
        elm.set(qn('w:space'), '4')
        elm.set(qn('w:color'), 'auto')
        pBdr.append(elm)
    pPr.append(pBdr)


def _insert_native_toc(after_paragraph):
    """
    在 after_paragraph 之后插入一个 Word 原生 TOC 域，
    支持更新标题和页码，涵盖级别 1–9。
    """
    # 构造 <w:p> 节点
    toc_p = OxmlElement('w:p')

    # 1) fldChar begin
    r1 = OxmlElement('w:r')
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    r1.append(fld_char_begin)
    toc_p.append(r1)

    # 2) instrText
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-9" \\h \\z \\u'
    r2.append(instr)
    toc_p.append(r2)

    # 3) fldChar separate
    r3 = OxmlElement('w:r')
    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    r3.append(fld_char_sep)
    toc_p.append(r3)

    # 4) 占位文本（可选）
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '右击此处更新目录'
    r4.append(t)
    toc_p.append(r4)

    # 5) fldChar end
    r5 = OxmlElement('w:r')
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    r5.append(fld_char_end)
    toc_p.append(r5)

    # 插入
    after_paragraph._p.addnext(toc_p)


def convert_ipynb_to_docx(ipynb_path, docx_path=None, page_size="A3"):
    """
    将 .ipynb 转为 .docx，并实现：
      1. 第一 Markdown 单元首行做文档标题
      2. 在第2行插入可更新的 Word 原生 TOC（1–9 级）
      3. 所有标题左对齐
      4. 仅给原始 code 单元段落加边框（不含输出）
      5. 表格等分列宽居中；图像放大至页宽并居中
      6. 若目标 docx 正被打开，抛出提示“请先关闭文件”
    """
    # 0. 强制保存当前 Notebook，貌似不管用，需要手动保存当前的Notebook
    #print("Saving current ipynb ...")
    #_save_current_notebook()
    
    # ---- 1. 检查输入 & 输出路径 ----
    if not os.path.isfile(ipynb_path):
        raise FileNotFoundError(f"找不到输入文件：{ipynb_path}")
    if docx_path is None:
        base, _ = os.path.splitext(ipynb_path)
        docx_path = base + ".docx"

    # ---- 2. 读取 Notebook，提取标题 & 收集 code 单元源码 ----
    nb = nbformat.read(ipynb_path, as_version=4)
    title = None
    code_blocks = []
    for cell in nb.cells:
        if cell.cell_type == "markdown" and title is None:
            lines = cell.source.strip().splitlines()
            if lines:
                title = lines[0].lstrip("# ").strip()
                # 去除这行，避免后面重复
                cell.source = "\n".join(lines[1:]).strip()
        if cell.cell_type == "code":
            code_blocks.append(cell.source.rstrip())
    if not title:
        title = os.path.splitext(os.path.basename(ipynb_path))[0]

    # ---- 3. 确保 Pandoc 可用 ----
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        pypandoc.download_pandoc()

    # ---- 4. Notebook → HTML（嵌入图像） ----
    exporter = HTMLExporter()
    exporter.embed_images = True

    # 关闭所有小于 CRITICAL 的日志
    logging.disable(logging.CRITICAL)

    buf = io.StringIO()
    # 屏蔽 stderr和stdout
    with contextlib.redirect_stderr(buf):
        with contextlib.redirect_stdout(buf):
            html_body, _ = exporter.from_notebook_node(nb)
    html = f"<h1>{title}</h1>\n" + html_body

    # 恢复日志
    #logging.disable(logging.NOTSET)

    # ---- 5. HTML → DOCX via Pandoc（或 subprocess fallback） ----
    try:
        pypandoc.convert_text(
            html, to="docx", format="html",
            outputfile=docx_path, encoding="utf-8"
        )
    except Exception:
        # fallback 到外部 pandoc
        with tempfile.NamedTemporaryFile("w", suffix=".html",
                                         delete=False,
                                         encoding="utf-8") as tmp:
            tmp.write(html)
            html_file = tmp.name
        try:
            try:
                subprocess.run(
                    ["pandoc", html_file, "-f", "html", "-t", "docx", "-o", docx_path],
                    check=True, capture_output=True
                )
            except subprocess.CalledProcessError as e:
                err = e.stderr.decode("utf-8", errors="ignore")
                low = err.lower()
                if "permission denied" in low or "could not open file" in low:
                    raise PermissionError(
                        f"无法写入 {docx_path}：文件可能已被打开，请先关闭后重试。"
                    )
                raise RuntimeError(f"Pandoc 转换失败：{err}")
        finally:
            os.remove(html_file)

    # ---- 6. 后处理 DOCX ----
    try:
        doc = Document(docx_path)
    except (PermissionError, OSError) as e:
        if getattr(e, "errno", None) in (errno.EACCES, errno.EPERM):
            raise PermissionError(
                f"无法打开 {docx_path}：文件可能已被打开，请先关闭后重试。"
            )
        raise

    # 6.1 页面尺寸 & 边距
    sec = doc.sections[0]
    if page_size.upper() in PAGE_SIZES:
        w_mm, h_mm = PAGE_SIZES[page_size.upper()]
        sec.page_width, sec.page_height = Mm(w_mm), Mm(h_mm)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Mm(25.4))
    avail_w = sec.page_width - sec.left_margin - sec.right_margin

    # 6.2 第一段替换为标题，设为 Heading1，左对齐
    p0 = doc.paragraphs[0]
    p0.text = title
    p0.style = doc.styles["Heading 1"]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 6.3 在第2行插入 Word 本地 TOC
    _insert_native_toc(p0)

    # 6.4 强制 Word 打开时自动更新目录
    try:
        settings = doc.settings.element
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)
    except Exception:
        pass  # 部分 python-docx 版本无此接口

    # 6.5 所有标题左对齐
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 6.6 仅给原始 code 单元对应段落加边框
    for p in doc.paragraphs:
        if "code" in p.style.name.lower():
            txt = p.text.rstrip()
            if any(txt == block for block in code_blocks):
                _add_border_to_paragraph(p)

    # 6.7 表格等分列宽并居中
    for tbl in doc.tables:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.allow_autofit = False
        cols = len(tbl.columns) or 1
        col_w = avail_w // cols
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = col_w

    # 6.8 图像放大至页宽并居中
    for shp in doc.inline_shapes:
        ow, oh = shp.width, shp.height
        fact = avail_w / ow
        shp.width = avail_w
        shp.height = int(oh * fact)
        p_el = shp._inline.getparent().getparent()
        Paragraph(p_el, doc).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 7. 保存并捕捉写入锁定 ----
    try:
        doc.save(docx_path)
    except (PermissionError, OSError) as e:
        if getattr(e, "errno", None) in (errno.EACCES, errno.EPERM):
            raise PermissionError(
                f"无法写入 {docx_path}：文件可能已被打开，请先关闭后重试。"
            )
        raise

    return docx_path


#==============================================================================
import os
import sys
import psutil

def is_file_opened(file_path: str) -> bool:
    """
    检测文件是否被其他程序打开（跨平台）
    :param file_path: 文件路径
    :return: True-被占用, False-未占用或不存在
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        return False
    
    abs_path = os.path.abspath(file_path)  # 转为绝对路径
    
    # 方法1：异常捕获法（快速检测）
    try:
        with open(abs_path, "a") as f:  # 追加模式（不破坏内容）
            pass
        return False  # 成功打开说明未被占用
    except (OSError, PermissionError):
        pass  # 继续尝试其他方法
    
    # 方法2：进程扫描法（精确检测）
    try:
        for proc in psutil.process_iter(['pid', 'name', 'open_files']):
            try:
                open_files = proc.info.get('open_files')
                if open_files and any(f.path == abs_path for f in open_files):
                    return True
            except (psutil.AccessDenied, psutil.NoSuchProcess):
                continue
    except NameError:  # psutil未安装
        pass
    
    # 方法3：文件锁试探法（最终回退）
    try:
        if sys.platform == 'win32':
            import msvcrt
            with open(abs_path, "a") as f:
                msvcrt.locking(f.fileno(), msvcrt.LK_NBLCK, 1)  # 非阻塞锁
        else:
            import fcntl
            with open(abs_path, "a") as f:
                fcntl.flock(f, fcntl.LOCK_EX | fcntl.LOCK_NB)  # 非阻塞独占锁
        return False
    except (OSError, BlockingIOError, ImportError):
        return True  # 所有检测均失败视为占用
    return False

#==============================================================================

def ipynb2docx(ipynb_path, page_size="A3"):
    """
    将 .ipynb 转为 .docx，特性：
      1. Markdown 首行做文档标题
      2. 在第 2 行插入全文 TOC（1–9 级）
      3. 所有标题左对齐，保留原字号
      4. 仅为“代码段”加边框，不影响输出
      5. 表格均分列宽并居中
      6. 图像放大至可用页宽并居中
      7. 若目标文件已打开，捕获并提示“请先关闭文件”
    """
    base, _ = os.path.splitext(ipynb_path)
    docx_path = base + ".docx"
    
    # 检测docx文件是否已被打开
    if is_file_opened(docx_path):
        print(f"Warning: {docx_path} occupied by other app, please close it and try again")
        return
    
    print(f"Converting to docx ...")
    
    result = convert_ipynb_to_docx(ipynb_path, docx_path=None, page_size=page_size)
    dirpath, filename = os.path.split(result)

    print(f"{filename} is created with TOC in {page_size} size")
    print(f"It is in {dirpath}")
    print(f"Note: may need further adjustments in formatting")
    
    return
    

