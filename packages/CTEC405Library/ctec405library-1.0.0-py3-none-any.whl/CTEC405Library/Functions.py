import os
import sys
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from PIL import Image
from keras.preprocessing.image import img_to_array
from nltk.stem import WordNetLemmatizer

stdoutInstance = sys.stdout

class _WordDocument():
    def __init__(self, file):
        self.file = file + ".docx";
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        style = self.document.styles['No Spacing']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        self.text = ''

    def writeText(self, text):
        if self.text == '':
            self.text = text
        else:
            self.text = self.text + text

    def saveDocument(self):
        paragraph = self.document.add_paragraph(self.text)
        paragraph.style = self.document.styles['No Spacing']
        self.document.save(self.file)

class _Logger(object):
    def __init__(self, file):
        self.terminal = sys.stdout
        self.wordDocument = _WordDocument(file)
   
    def write(self, message):
        self.terminal.write(message)
        self.wordDocument.writeText(message)

    def closeLog(self):
        self.wordDocument.saveDocument()

    def flush(self):
        # this flush method is needed for python 3 compatibility.
        # this handles the flush command by doing nothing.
        # you might want to specify some extra behavior here.
        pass

    def __del__(self):
        self.closeLog()
        sys.stdout = stdoutInstance

def enableLogging(file):
    sys.stdout = _Logger(file)

def disableLogging(file):
    sys.stdout.closeLog()
    sys.stdout = stdoutInstance
    
def printTable(array, columnNames=None):
    arrayShape = array.shape
    if len(arrayShape) == 1:
        array = np.reshape(array, (1, arrayShape[0]))
        arrayShape = array.shape
    if columnNames is None:
        if len(arrayShape) == 1:
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            df = pd.DataFrame(array, index=blankRows, columns=[''])
        elif len(arrayShape) > 1:
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            numColumns = int(arrayShape[1])
            blankColumns = ['' for i in range(numColumns)]
            df = pd.DataFrame(array, index=blankRows, columns=blankColumns)
    else:
        numRows = int(arrayShape[0])
        blankRows = ['' for i in range(numRows)]
        df = pd.DataFrame(array, index=blankRows, columns=columnNames)
    pd.set_option("display.max_rows", None)
    pd.set_option("display.max_columns", None)
    pd.set_option('display.expand_frame_repr', False)
    print(df)
    print("")
    
def printVariable(label, var, columnNames=None):
    if type(var) is np.ndarray:
        print(label + ":")
        printTable(var, columnNames)
    elif type(var) is list:
        print(label + ":")
        for item in var:
            print("\t", item)
        print("")
    else:
        print(label + ": " + str(var))

def readExcelSpreadsheet(file, useColumnNames=True, useRowLabels=False):
    if useColumnNames:
        headerIndex = 0
    else:
        headerIndex = None

    if useRowLabels:
        rowIndex = 0
    else:
        rowIndex = None
    data = pd.read_excel(file, header=headerIndex, index_col=rowIndex)
    dataNumbers = data.to_numpy()
    indexList = list(data.index)
    namesList = list(data.columns)

    return dataNumbers, indexList, namesList

def encodeLabels(YRaw):
  data = pd.DataFrame(YRaw)
  indexList = list(data[data.columns[0]])
  uniqueIndexesList = list(data[data.columns[0]].unique())
  firstCategory = indexList[0]
  try:
      int(firstCategory)
      YIdx = [uniqueIndexesList.index(idx) for idx in uniqueIndexesList]
      YData = [int(idx) for idx in uniqueIndexesList]
      YDiff = set(YIdx).difference(YData)
      if len(YDiff) == 0:
          Y = np.array([int(idx) for idx in indexList])
      else:
          Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
  except ValueError:
      Y = np.array([uniqueIndexesList.index(idx) for idx in indexList])
      namesList = list(data.columns)
  return Y

def readImages(imageData, imageDir, dataCol=0, categoriesCol=-1):
    imageArrayList = []
    imagePathList = [imageDir + im[dataCol] for im in imageData]
    for i in range(len(imagePathList)):
        imagePath = imagePathList[i]
        img = Image.open(imagePath).convert("RGB")
        category = imageData[i, categoriesCol]
        imageArray = img_to_array(img).flatten()
        imageArray = np.hstack([imageArray, category])
        imageArrayList.append(imageArray)

    return np.array(imageArrayList)

def convertArray(var):
  if type(var) is np.ndarray:
    return np.array_str(var)
  return str(var)

def convertList(var):
  if type(var) is list:
    listString = ' '.join(map(str, var))
    return listString
  return str(var)

def getPrompt(*args):
  prompt = ''
  for i in range(len(args)):
    arg = args[i]
    if i < len(args) - 1:
      if type(arg) is np.ndarray:
        prompt = prompt + np.array_str(arg) + ' '
      elif type(arg) is list:
        prompt = prompt + convertList(arg) + ' '
      else:
        prompt = prompt + str(arg) + ' '
    else:
      if type(arg) is np.ndarray:
        prompt = prompt + np.array_str(arg)
      elif type(arg) is list:
        prompt = prompt + convertList(arg)
      else:
        prompt = prompt + str(arg)
  return prompt

def getRootWords(words):
    if type(words) is list:
        wordList = words
    elif type(words) is np.ndarray:
        wordList = words.flatten().tolist()
    wnl = WordNetLemmatizer()
    lemmaWords = []
    for word in wordList:
        word = word.lower()
        lemmaWord = wnl.lemmatize(word)
        lemmaWords.append(lemmaWord)
    lemmaWordsArray = np.array(lemmaWords)
    uniqueWordsArray = np.unique(lemmaWordsArray)
    return uniqueWordsArray
      
  

